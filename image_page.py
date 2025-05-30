from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_image_page(document, images, title, caption_prefix="Figure", description=None):
    """Add a dedicated page for images with improved layout and custom captions."""

    # Add title
    heading = document.add_paragraph()
    run = heading.add_run(title.upper())
    run.font.size = Pt(16)
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add description if provided
    if description:
        desc_para = document.add_paragraph(description)
        desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        document.add_paragraph()  # space after description

    # Add images 2 per row, with captions directly below each
    for i in range(0, len(images), 2):
        # First image and caption
        if i < len(images):
            try:
                img1_para = document.add_paragraph()
                run1 = img1_para.add_run()
                run1.add_picture(images[i], width=Inches(4.5))
                img1_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                caption1 = document.add_paragraph(f"{caption_prefix} {i+1}")
                caption1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                caption1.style = 'Caption'
            except Exception as e:
                document.add_paragraph(f"Could not load image: {os.path.basename(images[i])}")

        # Second image and caption
        if i+1 < len(images):
            try:
                img2_para = document.add_paragraph()
                run2 = img2_para.add_run()
                run2.add_picture(images[i+1], width=Inches(4.5))
                img2_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                caption2 = document.add_paragraph(f"{caption_prefix} {i+2}")
                caption2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                caption2.style = 'Caption'
            except Exception as e:
                document.add_paragraph(f"Could not load image: {os.path.basename(images[i+1])}")

        # Add vertical space between rows
        document.add_paragraph()

