import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}

def fetch_objective(topic, report_type='college'):
    """
    Fetch Objective content for the given topic using the Gemini API.
    The response will be formatted differently for school vs college reports.
    """
    if report_type == 'school':
        prompt = (
            f"Generate professional objectives for a {topic} project with this format:\n\n"
            f"The main goal of a {topic} is to make . Here are the main things it aims to do:\n\n"
            "1. [Technical Objective 1]: [Detailed description]\n"
            "2. [Technical Objective 2]: [Detailed description]\n"
            "3. [Technical Objective 3]: [Detailed description]\n"
            "4. [Technical Objective 4]: [Detailed description]\n"
            "5. [Technical Objective 5]: [Detailed description]\n\n"
            "Example for Hospital Management System:\n"
            "1. Patient Management: Allow hospitals to manage patient records, appointments, and medical history efficiently. This includes storing personal details, tracking medical treatments, and ensuring data security, all of which contribute to streamlined patient care.\n"
            "2. Appointment Scheduling: Enable patients and doctors to schedule, reschedule, or cancel appointments easily. Patients can book, reschedule, or cancel appointments online, while doctors can maintain their availability calendar. \n"
            "3. Billing and Payments: Automate billing processes and provide payment options for patients. It also maintains historical billing records for transparency and ease of reference.\n"
            "4. Inventory Management: Track and manage medical supplies, equipment, and pharmaceuticals.  The system should monitor stock levels in real-time, trigger low-stock alerts, manage vendor information, and automate the generation of purchase orders to ensure uninterrupted medical services.\n"
            "5. Report Generation: Generate reports for patient statistics, financial records, and inventory status. Features such as role-based dashboards, responsive design for mobile compatibility, multilingual support, and clear navigation ensure that users can interact with the system effectively, regardless of their technical proficiency.\n"
            "6. User-Friendly Interface: Provide an intuitive interface for doctors, staff, and patients to interact with the system.\n\n"
            f"Now create objectives for {topic}:"
        )
    else:
        prompt = (
            f"Generate professional objectives for a {topic} project with this format:\n\n"
            f"The main goal of a {topic} is to make . Here are the main things it aims to do:\n\n"
            "1. [Technical Objective 1]: [Detailed description]\n"
            "2. [Technical Objective 2]: [Detailed description]\n"
            "3. [Technical Objective 3]: [Detailed description]\n"
            "4. [Technical Objective 4]: [Detailed description]\n"
            "5. [Technical Objective 5]: [Detailed description]\n\n"
            "Example for Hospital Management System:\n"
            "1. Patient Management: Allow hospitals to manage patient records, appointments, and medical history efficiently. This includes storing personal details, tracking medical treatments, and ensuring data security, all of which contribute to streamlined patient care.\n"
            "2. Appointment Scheduling: Enable patients and doctors to schedule, reschedule, or cancel appointments easily. Patients can book, reschedule, or cancel appointments online, while doctors can maintain their availability calendar. \n"
            "3. Billing and Payments: Automate billing processes and provide payment options for patients. It also maintains historical billing records for transparency and ease of reference.\n"
            "4. Inventory Management: Track and manage medical supplies, equipment, and pharmaceuticals.  The system should monitor stock levels in real-time, trigger low-stock alerts, manage vendor information, and automate the generation of purchase orders to ensure uninterrupted medical services.\n"
            "5. Report Generation: Generate reports for patient statistics, financial records, and inventory status. Features such as role-based dashboards, responsive design for mobile compatibility, multilingual support, and clear navigation ensure that users can interact with the system effectively, regardless of their technical proficiency.\n"
            "6. User-Friendly Interface: Provide an intuitive interface for doctors, staff, and patients to interact with the system.\n\n"
            f"Now create objectives for {topic}:"
        )

    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }
    
    params = {'key': GEMINI_API_KEY}

    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()
        
        if 'candidates' in content and len(content['candidates']) > 0:
            return content['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"No objectives generated for {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching Objective content: {e}"

def post_process_objective_to_doc(text, document, report_type='college'):
    """
    Post-process the generated objective text and add it to the Word document.
    Applies different formatting for school vs college reports.
    """
    # Set Times New Roman as Default Font
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Clean and split text
    text = text.replace("**", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # School report formatting
    if report_type == 'school':
        for line in lines:
            para = document.add_paragraph(line)
            # para.paragraph_format.space_after = Pt(12)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Improved justification control
            para.paragraph_format.widow_control = True  # Prevent single words on last line
            para.paragraph_format.keep_together = True  # Keep paragraph together
            para.paragraph_format.keep_with_next = True
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            # Adjust hyphenation (helps with justification)
            para.paragraph_format.word_wrap = True
            para.paragraph_format.hyphenation = True
    
    # College report formatting
    else:
        intro_processed = False
        previous_line = None

        for line in lines:
            if line == previous_line:
                continue

            if not intro_processed:
                para = document.add_paragraph(line)
                # para.paragraph_format.space_after = Pt(21)
                intro_processed = True
                previous_line = line

                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue

            if ":" in line and not line.startswith("Feasibility Study"):
                heading, desc = line.split(":", 1)
                para = document.add_paragraph()
                run_heading = para.add_run(f"{heading}:")
                run_heading.bold = True
                run_heading.underline = True
                run_heading.font.size = Pt(12)
                
                run_desc = para.add_run(f" {desc.strip()}")
                run_desc.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True  # Prevent single words on last line
                para.paragraph_format.keep_together = True  # Keep paragraph together
                para.paragraph_format.keep_with_next = True

                # para.paragraph_format.space_after = Pt(19)

                # Adjust hyphenation (helps with justification)
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True

            else:
                para = document.add_paragraph(line)
                # para.paragraph_format.space_after = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True  # Prevent single words on last line
                para.paragraph_format.keep_together = True  # Keep paragraph together
                para.paragraph_format.keep_with_next = True
                
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

                # Adjust hyphenation (helps with justification)
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True