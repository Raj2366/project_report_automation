import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}

def fetch_modules_and_descriptions(topic, report_type='college'):
    """
    Fetch Modules and Descriptions content for the given topic using the Gemini API.
    The response will be formatted with numbered modules and detailed descriptions.
    """
    if report_type == 'school':
        prompt = (
            f"Generate professional modules and descriptions for a {topic} project with this format:\n\n"
            
            "1. [Module Name 1]\n"
            "Description: [Detailed description of the module's purpose and functionality]\n\n"
            "2. [Module Name 2]\n"
            "Description: [Detailed description of the module's purpose and functionality]\n\n"
            "3. [Module Name 3]\n"
            "Description: [Detailed description of the module's purpose and functionality]\n\n"
            "4. [Module Name 4]\n"
            "Description: [Detailed description of the module's purpose and functionality]\n\n"
            "Example for Hospital Management System:\n"
            "1. Patient Registration Module\n"
            "Description: Handles the registration of new patients and maintenance of patient demographics. "
            "It captures personal details, insurance information, and medical history, creating a unique "
            "identifier for each patient to ensure accurate record-keeping across all hospital departments.\n\n"
            "2. Appointment Management Module\n"
            "Description: Manages scheduling, rescheduling, and cancellation of patient appointments. "
            "It provides real-time availability of doctors, sends automated reminders to patients, "
            "and maintains a calendar view for staff to optimize resource allocation.\n\n"
            "3. Billing and Insurance Module\n"
            "Description: Automates the generation of medical bills, processes insurance claims, "
            "and tracks payments. It integrates with treatment records to ensure accurate billing, "
            "supports multiple payment methods, and generates financial reports for accounting purposes.\n\n"
            "4. Pharmacy and Inventory Module\n"
            "Description: Manages medication dispensing, tracks drug inventory levels, and automates "
            "reordering processes. It validates prescriptions against patient records, maintains drug "
            "interaction databases, and generates alerts for expiring medications.\n\n"
            f"Now create modules and descriptions for {topic}:"
        )
    else:
        prompt = (
            f"Generate professional modules and descriptions for a {topic} project with this format:\n\n"
            
            "1. [Module Name 1]\n"
            "Description: [Detailed description of the module's purpose and functionality]\n\n"
            "2. [Module Name 2]\n"
            "Description: [Detailed description of the module's purpose and functionality]\n\n"
            "3. [Module Name 3]\n"
            "Description: [Detailed description of the module's purpose and functionality]\n\n"
            "4. [Module Name 4]\n"
            "Description: [Detailed description of the module's purpose and functionality]\n\n"
            "Example for Hospital Management System:\n"
            "1. Patient Registration Module\n"
            "Description: Handles the registration of new patients and maintenance of patient demographics. "
            "It captures personal details, insurance information, and medical history, creating a unique "
            "identifier for each patient to ensure accurate record-keeping across all hospital departments.\n\n"
            "2. Appointment Management Module\n"
            "Description: Manages scheduling, rescheduling, and cancellation of patient appointments. "
            "It provides real-time availability of doctors, sends automated reminders to patients, "
            "and maintains a calendar view for staff to optimize resource allocation.\n\n"
            "3. Billing and Insurance Module\n"
            "Description: Automates the generation of medical bills, processes insurance claims, "
            "and tracks payments. It integrates with treatment records to ensure accurate billing, "
            "supports multiple payment methods, and generates financial reports for accounting purposes.\n\n"
            "4. Pharmacy and Inventory Module\n"
            "Description: Manages medication dispensing, tracks drug inventory levels, and automates "
            "reordering processes. It validates prescriptions against patient records, maintains drug "
            "interaction databases, and generates alerts for expiring medications.\n\n"
            f"Now create modules and descriptions for {topic}:"
        )

    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }
    
    params = {'key': GEMINI_API_KEY}

    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()
        
        if 'candidates' in content and len(content['candidates']) > 0:
            return content['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"No modules and descriptions generated for {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching Modules and Descriptions content: {e}"

def post_process_modules_to_doc(text, document, report_type='college'):
    """
    Post-process the generated modules text and add it to the Word document.
    Applies different formatting for school vs college reports.
    """
    # Set Times New Roman as Default Font
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Clean and split text
    text = text.replace("**", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # School report formatting
    if report_type == 'school':
        intro_processed = False
        previous_line = None

        for line in lines:
            if line == previous_line:
                continue

            if not intro_processed:
                para = document.add_paragraph(line)
                intro_processed = True
                previous_line = line

                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue

            if ":" in line:
                heading, desc = line.split(":", 1)
                para = document.add_paragraph()


                if "[" in heading and "]" in heading:
                    start = heading.index("[")
                    end = heading.index("]") + 1
                    before = heading[:start].strip()
                    bold_part = heading[start:end].strip()
                    after = heading[end:].strip()

                    if before:
                        para.add_run(before+" ").font.size = Pt(12)

                    bold_run = para.add_run(bold_part)
                    bold_run.bold = True
                    bold_run.font.size = Pt(12)

                    if after:
                        para.add_run(" "+after + ":").font.size = Pt(12)
                    else:
                        para.add_run(":").font.size = Pt(12)
                else:
                    run_heading = para.add_run(f"{heading.strip()}:")
                    run_heading.bold = True
                    run_heading.font.size = Pt(12)

                
                run_desc = para.add_run(f" {desc.strip()}")
                run_desc.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_together = True
                para.paragraph_format.keep_with_next = True

                # Adjust hyphenation
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True

            else:
                para = document.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_together = True
                para.paragraph_format.keep_with_next = True
                
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

                # Adjust hyphenation
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True
    
    # College report formatting
    else:
        intro_processed = False
        previous_line = None

        for line in lines:
            if line == previous_line:
                continue

            if not intro_processed:
                para = document.add_paragraph(line)
                intro_processed = True
                previous_line = line

                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue

            if ":" in line:
                heading, desc = line.split(":", 1)
                para = document.add_paragraph()

                if "[" in heading and "]" in heading:
                    start = heading.index("[")
                    end = heading.index("]") + 1
                    before = heading[:start].strip()
                    bold_part = heading[start:end].strip()
                    after = heading[end:].strip()

                    if before:
                        para.add_run(before+" ").font.size = Pt(12)

                    bold_run = para.add_run(bold_part)
                    bold_run.bold = True
                    bold_run.font.size = Pt(12)

                    if after:
                        para.add_run(" "+after + ":").font.size = Pt(12)
                    else:
                        para.add_run(":").font.size = Pt(12)
                else:
                    run_heading = para.add_run(f"{heading.strip()}:")
                    run_heading.bold = True
                    run_heading.font.size = Pt(12)

                
                run_desc = para.add_run(f" {desc.strip()}")
                run_desc.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_together = True
                para.paragraph_format.keep_with_next = True

                # Adjust hyphenation
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True

            else:
                para = document.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_together = True
                para.paragraph_format.keep_with_next = True
                
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

                # Adjust hyphenation
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True