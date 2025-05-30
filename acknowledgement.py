from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def generate_acknowledgement_page(document, submitted_by, submitted_to, department, institute_name, **kwargs):

    ack_text = (
        f"I would like to express a deep sense of gratitude and thanks "
        f"profusely to {submitted_to} without his/her wise counsel and able guidance, it would have been "
        f"impossible to complete the project in this manner. I express gratitude to other faculty members of "
        f"{department} department of {institute_name} for their intellectual support throughout the course "
        f"of this work. Finally, I am indebted to all whosoever have contributed in this report work."
    )

    paragraphs = ack_text.split('\n')
    for para in paragraphs:
        if para.strip():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.widow_control = True
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = True

            run = p.add_run(para.strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # Name paragraph
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(f"\nName: {submitted_by}")
    name_run.font.name = 'Times New Roman'
    name_run.font.size = Pt(12)

    # Date paragraph
    date_paragraph = document.add_paragraph()
    date_run = date_paragraph.add_run("Date: " + datetime.today().strftime("%d-%m-%Y"))
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(12)
