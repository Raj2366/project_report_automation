import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}

def fetch_unique_features(topic, report_type='college'):
    """
    Fetch Unique Features content for the given topic using the Gemini API.
    The response will be formatted differently for school vs college reports.
    """
    if report_type == 'school':
        prompt = (
            f"Generate professional unique features for a {topic} project with this format:\n\n"
        f"1. [Feature Name]:"
        "[Detailed description]\n\n"
        "2. [Feature Name]:"
        "[Detailed description]\n\n"
        "3. [Feature Name]:"
        "[Detailed description]\n\n"
        "4. [Feature Name]:"
        "[Detailed description]\n\n"
        "Example for Personal Finance Tracker:"
        "1. Seamless Multi-Device Synchronization:"
        "Utilizing SQLite's lightweight database capabilities, the system allows for efficient synchronization across multiple devices. Users can access and update their financial data from their desktop, laptop, or mobile devices without any discrepancies.\n\n"
        "2. Customizable Budgeting Tools:"
        "Users can create, manage, and customize budgets tailored to their specific financial goals. The system provides flexibility in setting budget limits for various expense categories, helping users stay on track with their financial plans.\n\n"
        "3. Advanced Data Analysis and Reporting:"
        "The system offers comprehensive data analysis features, including detailed financial reports, spending summaries, and visual analytics such as charts and graphs. This allows users to gain insights into their spending patterns and make informed financial decisions.\n\n"
        "4. Report Generation:"
        "Generate reports for patient statistics, financial records, and inventory status. Features such as role-based dashboards, responsive design for mobile compatibility, multilingual support, and clear navigation ensure that users can interact with the system effectively, regardless of their technical proficiency.\n\n"
        f"Now create unique features for {topic}:"
        )
    else:
        prompt = (
            f"Generate professional unique features for a {topic} project with this format:\n\n"
        f"1. [Feature Name]:\n"
        "[Detailed description]\n\n"
        "2. [Feature Name]:\n"
        "[Detailed description]\n\n"
        "3. [Feature Name]:\n"
        "[Detailed description]\n\n"
        "4. [Feature Name]:\n"
        "[Detailed description]\n\n"
        "Example for Personal Finance Tracker:\n"
        "1. Seamless Multi-Device Synchronization:\n"
        "Utilizing SQLite's lightweight database capabilities, the system allows for efficient synchronization across multiple devices. Users can access and update their financial data from their desktop, laptop, or mobile devices without any discrepancies.\n\n"
        "2. Customizable Budgeting Tools:\n"
        "Users can create, manage, and customize budgets tailored to their specific financial goals. The system provides flexibility in setting budget limits for various expense categories, helping users stay on track with their financial plans.\n\n"
        "3. Advanced Data Analysis and Reporting:\n"
        "The system offers comprehensive data analysis features, including detailed financial reports, spending summaries, and visual analytics such as charts and graphs. This allows users to gain insights into their spending patterns and make informed financial decisions.\n\n"
        f"Now create unique features for {topic}:"
        )


    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }

    params = {'key': GEMINI_API_KEY}

    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()

        if 'candidates' in content and len(content['candidates']) > 0:
            return content['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"No unique features generated for {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching Unique Features content: {e}"

def post_process_unique_features_to_doc(text, document, report_type='college'):
    """
    Post-process the generated unique features text and add it to the Word document.
    Applies different formatting for school vs college reports.
    """
    # Set Times New Roman as Default Font
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Clean and split text
    text = text.replace("**", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # School report formatting
    if report_type == 'school':
        for line in lines:
            para = document.add_paragraph(line)
            para.paragraph_format.space_after = Pt(12)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    
    # College report formatting
    else:
        # intro_processed = False
        previous_line = None

        for line in lines:
            if line == previous_line:
                continue

            if ":" in line and not line.startswith("Feasibility Study"):
                heading, desc = line.split(":", 1)
                para = document.add_paragraph()
                run_heading = para.add_run(f"{heading}:")
                run_heading.bold = True
                run_heading.underline = True
                run_heading.font.size = Pt(12)
                
                run_desc = para.add_run(f" {desc.strip()}")
                run_desc.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(19)

            else:
                para = document.add_paragraph(line)
                para.paragraph_format.space_after = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)