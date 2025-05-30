import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}

def fetch_limitations(topic, report_type='college'):
    """
    Fetch Limitations content for the given topic using the Gemini API.
    The response will be formatted with bullet points for issues and solutions.
    """
    if report_type == 'school':
        prompt = (
            f"Generate limitations for a {topic} project with this format:\n\n"
            f"The {topic} has these current limitations:\n\n"
            "[Limitation Name]\n"
            "• Issue: [Description of the limitation]\n"
            "• Solution: [Suggested solution to overcome this limitation]\n\n"
            "[Limitation Name]\n"
            "• Issue: [Description of the limitation]\n"
            "• Solution: [Suggested solution to overcome this limitation]\n\n"
            "Example for Hospital Management System:\n"
            "Limited Interoperability\n"
            "• Issue: The system may not integrate well with other hospital software, causing data silos.\n"
            "• Solution: Implement standard HL7/FHIR protocols for better healthcare data exchange.\n\n"
            "Mobile Accessibility\n"
            "• Issue: The system may not have full functionality on mobile devices, limiting doctor mobility.\n"
            "• Solution: Develop responsive web design or dedicated mobile apps for all critical functions.\n\n"
            f"Now create limitations for {topic}:"
        )
    else:
        prompt = (
            f"Generate technical limitations for a {topic} project with this format:\n\n"
            f"The {topic} has these technical constraints:\n\n"
            "[Technical Limitation Name]\n"
            "• Issue: [Technical description of the limitation]\n"
            "• Solution: [Technical solution approach to address this limitation]\n\n"
            "[Technical Limitation Name]\n"
            "• Issue: [Technical description of the limitation]\n"
            "• Solution: [Technical solution approach to address this limitation]\n\n"
            "Example for Hospital Management System:\n"
            "Data Processing Bottlenecks\n"
            "• Issue: The system experiences performance degradation when processing large medical imaging datasets due to sequential processing architecture.\n"
            "• Solution: Implement distributed computing using Apache Spark and GPU acceleration for parallel image processing tasks.\n\n"
            "Legacy System Integration\n"
            "• Issue: Difficulty integrating with older hospital information systems that use proprietary data formats and outdated protocols.\n"
            "• Solution: Develop middleware adapters using Python's Django framework with custom API gateways to translate between systems.\n\n"
            f"Now create limitations for {topic}:"
        )

    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }
    
    params = {'key': GEMINI_API_KEY}

    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()
        
        if 'candidates' in content and len(content['candidates']) > 0:
            return content['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"No limitations generated for {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching Limitations content: {e}"

def post_process_limitations_to_doc(text, document, report_type='college'):
    """
    Post-process the generated limitations text and add it to the Word document.
    Applies different formatting for school vs college reports.
    """
    # Set Times New Roman as Default Font
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Clean and split text
    text = text.replace("**", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # School report formatting
    if report_type == 'school':
        for line in lines:
            para = document.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Improved justification control
            para.paragraph_format.widow_control = True
            para.paragraph_format.keep_together = True
            para.paragraph_format.keep_with_next = True
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            # Adjust hyphenation
            para.paragraph_format.word_wrap = True
            para.paragraph_format.hyphenation = True
    
    # College report formatting
    else:
        intro_processed = False
        previous_line = None

        for line in lines:
            if line == previous_line:
                continue

            if not intro_processed:
                para = document.add_paragraph(line)
                intro_processed = True
                previous_line = line

                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue

            if "• Issue:" in line:
                para = document.add_paragraph(style='List Bullet')
                run = para.add_run(line.replace("• Issue:", "").strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Formatting controls
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_together = True
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True

            elif "• Solution:" in line:
                para = document.add_paragraph(style='List Bullet')
                run = para.add_run(line.replace("• Solution:", "").strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Formatting controls
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_together = True
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True

            else:
                para = document.add_paragraph()
                run = para.add_run(line)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Formatting controls
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_together = True
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True