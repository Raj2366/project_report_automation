import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}


def fetch_future_scope(topic, report_type='college'):
    """
    Fetch Future Scope content for the given topic using the Gemini API.
    The response will be formatted with bullet points and detailed descriptions.
    """
    if report_type == 'school':
        prompt = (
            f"Generate technical future enhancements for a {topic} project with this format:\n\n"
            f"The {topic} has these potential for these future improvements:\n\n"
            "1. [Advanced Feature Name 1]: [Detailed technical description of implementation]\n"
            "2. [Advanced Feature Name 2]: [Detailed technical description of implementation]\n"
            "3. [Advanced Feature Name 3]: [Detailed technical description of implementation]\n"
            "4. [Advanced Feature Name 4]: [Detailed technical description of implementation]\n\n"
            "Example for Hospital Management System:\n"
            "1. Predictive Analytics Engine: Develop ML models using Python/TensorFlow to predict admission rates and medication needs based on historical data patterns.\n"
            "2. Blockchain Medical Records: Implement Ethereum-based decentralized ledger for tamper-proof records with patient-controlled access permissions.\n"
            "3. IoT Smart Hospital: Network of sensors to monitor equipment status and patient movements in real-time with automated maintenance alerts.\n"
            "4. AR Surgical Assistance: Augmented reality system to overlay patient vitals and imaging data during surgical procedures for precision.\n\n"
            f"Now create future scope for {topic}:"
        )
    else:
        prompt = (
            f"Generate technical future enhancements for a {topic} project with this format:\n\n"
            f"The {topic} has these potential for these future improvements:\n\n"
            "1. [Advanced Feature Name 1]: [Detailed technical description of implementation]\n"
            "2. [Advanced Feature Name 2]: [Detailed technical description of implementation]\n"
            "3. [Advanced Feature Name 3]: [Detailed technical description of implementation]\n"
            "4. [Advanced Feature Name 4]: [Detailed technical description of implementation]\n\n"
            "Example for Hospital Management System:\n"
            "1. Predictive Analytics Engine: Develop ML models using Python/TensorFlow to predict admission rates and medication needs based on historical data patterns.\n"
            "2. Blockchain Medical Records: Implement Ethereum-based decentralized ledger for tamper-proof records with patient-controlled access permissions.\n"
            "3. IoT Smart Hospital: Network of sensors to monitor equipment status and patient movements in real-time with automated maintenance alerts.\n"
            "4. AR Surgical Assistance: Augmented reality system to overlay patient vitals and imaging data during surgical procedures for precision.\n\n"
            f"Now create future scope for {topic}:"
        )

    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }
    
    params = {'key': GEMINI_API_KEY}

    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()
        
        if 'candidates' in content and len(content['candidates']) > 0:
            return content['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"No future scope generated for {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching Future Scope content: {e}"

def post_process_future_scope_to_doc(text, document, report_type='college'):
    """
    Post-process the generated future scope text and add it to the Word document.
    Applies different formatting for school vs college reports.
    """
    # Set Times New Roman as Default Font
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Clean and split text
    text = text.replace("**", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # School report formatting
    if report_type == 'school':
        for line in lines:
            para = document.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Improved justification control
            para.paragraph_format.widow_control = True
            para.paragraph_format.keep_together = True
            para.paragraph_format.keep_with_next = True
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            # Adjust hyphenation
            para.paragraph_format.word_wrap = True
            para.paragraph_format.hyphenation = True
    
    # College report formatting
    else:
        intro_processed = False
        previous_line = None

        for line in lines:
            if line == previous_line:
                continue

            if not intro_processed:
                para = document.add_paragraph(line)
                intro_processed = True
                previous_line = line

                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue

            if ":" in line:
                heading, desc = line.split(":", 1)
                para = document.add_paragraph()
                run_heading = para.add_run(f"{heading}:")
                run_heading.bold = True
                run_heading.underline = True
                run_heading.font.size = Pt(12)
                
                run_desc = para.add_run(f" {desc.strip()}")
                run_desc.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_together = True
                para.paragraph_format.keep_with_next = True

                # Adjust hyphenation
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True

            else:
                para = document.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_together = True
                para.paragraph_format.keep_with_next = True
                
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

                # Adjust hyphenation
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True