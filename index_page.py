from docx import Document
from docx.shared import Pt, Cm
import tempfile
import os
import subprocess
import logging
import shutil
from typing import Dict, List, Optional, Tuple
import re

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class IndexGenerator:
    """Generates index with accurate page ranges using PDF analysis"""
    
    def __init__(self):
        # Mapping of section names to display names
        self.section_map = {
            "Introduction": "Introduction",
            "Objective": "Objective",
            "Problem Formulation": "Problem Formulation",
            "Feasibility": "Feasibility Study of Project",
            "DFD": "Data Flow Diagram",
            "Flowchart": "Flowchart",
            "Future Scope": "Future Scope",
            "Conclusion": "Conclusion",
            "Unique Features": "Unique Features",
            "Reference": "References/Bibliography",
            "Modules and Descriptions": "Modules and Descriptions",
            "Limitations": "Limitations",
            "Abstract": "Abstract",
            "Project Images": "Project Snapshots",
            "Code Screenshots": "Code Snapshots"
        }
        
        # Try to locate LibreOffice automatically
        self.libreoffice_path = self._find_libreoffice()

    def _find_libreoffice(self) -> Optional[str]:
        """Try to locate LibreOffice executable"""
        paths = [
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
            "/usr/bin/libreoffice",
            "/usr/local/bin/libreoffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        ]
        
        for path in paths:
            if os.path.exists(path):
                return path
        return None

    def create_index_page(
        self,
        document: Document,
        ordered_sections: List[str],
        custom_sections: List[str],
        full_doc_path: str
    ) -> None:
        """
        Create index page with accurate page ranges
        Args:
            document: Target document to add index to
            ordered_sections: List of sections in order
            custom_sections: List of custom sections
            full_doc_path: Path to complete document for page number reference
        """
        try:
            # Get page ranges from complete document
            heading_ranges = self._get_page_ranges_from_complete_doc(full_doc_path)
            
            # Build the index table with page ranges
            self._build_index_table(document, ordered_sections, custom_sections, heading_ranges)
            
        except Exception as e:
            logger.error(f"Error generating index: {e}")
            # Fallback with empty page numbers
            self._build_index_table(document, ordered_sections, custom_sections, {})

    def _get_page_ranges_from_complete_doc(self, docx_path: str) -> Dict[str, Tuple[int, int]]:
        """Generate PDF from complete doc and extract page ranges"""
        with tempfile.TemporaryDirectory() as temp_dir:
            # Convert complete document to PDF
            pdf_path = os.path.join(temp_dir, "complete.pdf")
            self._convert_to_pdf(docx_path, pdf_path)
            
            # Extract heading ranges from PDF
            return self._extract_pdf_heading_ranges(pdf_path)

    def _convert_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        """Convert DOCX to PDF using LibreOffice"""
        try:
            cmd = [
                self.libreoffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(pdf_path),
                docx_path
            ]
            
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=30,
                text=True
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
            
            if not os.path.exists(pdf_path):
                raise RuntimeError("PDF output not found after conversion")
                
        except Exception as e:
            raise RuntimeError(f"PDF conversion error: {e}")

    def _extract_pdf_heading_ranges(self, pdf_path: str) -> Dict[str, Tuple[int, int]]:
        """
        Extract headings and their page ranges from PDF
        Returns: Dictionary of {heading_text: (start_page, end_page)}
        """
        try:
            from pdfminer.high_level import extract_pages
            from pdfminer.layout import LTTextContainer, LTTextBoxHorizontal
        except ImportError:
            logger.warning("pdfminer.six not installed - cannot extract PDF headings")
            return {}

        heading_ranges = {}
        try:
            current_heading = None
            start_page = 1
            previous_page_headings = []
            
            for page_num, page_layout in enumerate(extract_pages(pdf_path), 1):
                page_headings = []
                
                # Find all headings on current page
                for element in page_layout:
                    if isinstance(element, LTTextBoxHorizontal):
                        text = element.get_text().strip()
                        if text:
                            # Clean and normalize the text
                            clean_text = ' '.join(text.split())
                            clean_text = clean_text.replace('\n', ' ').replace('  ', ' ')
                            
                            # Skip page numbers and small text (likely not headings)
                            if (not clean_text.isdigit() and 
                                len(clean_text) < 100 and 
                                len(clean_text.split()) < 10 and
                                not clean_text.endswith('%') and
                                not re.match(r'^[0-9.,]+$', clean_text)):
                                
                                # Further clean the text by removing special characters
                                clean_text = re.sub(r'[^a-zA-Z0-9 \-\n]', '', clean_text)
                                page_headings.append(clean_text)
                
                # If we found new headings that weren't on the previous page
                new_headings = [h for h in page_headings if h not in previous_page_headings]
                
                if new_headings:
                    # Finalize previous heading range
                    if current_heading:
                        heading_ranges[current_heading] = (start_page, page_num - 1)
                    
                    # Start new range with the most significant heading
                    current_heading = self._select_primary_heading(new_headings)
                    start_page = page_num
                
                previous_page_headings = page_headings
            
            # Finalize the last heading range
            if current_heading and current_heading not in heading_ranges:
                heading_ranges[current_heading] = (start_page, page_num)
                
            return heading_ranges
            
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            return {}

    def _select_primary_heading(self, headings: List[str]) -> str:
        """
        Select the most significant heading from a list of candidates
        """
        # Prioritize longer headings that match our known sections
        for heading in sorted(headings, key=len, reverse=True):
            normalized = heading.lower().strip()
            for section in self.section_map.values():
                if section.lower() in normalized:
                    return section
        return headings[0] if headings else ""

    def _build_index_table(
        self,
        document: Document,
        ordered_sections: List[str],
        custom_sections: List[str],
        heading_ranges: Dict[str, Tuple[int, int]]
    ) -> None:
        """Build the index table with page ranges"""
        # Add index heading
        document.add_page_break()
        heading = document.add_paragraph("INDEX")
        heading.runs[0].font.size = Pt(18)
        heading.runs[0].bold = True
        heading.runs[0].underline = True
        heading.runs[0].font.name = "Times New Roman"
        heading.alignment = 1  # Center alignment
        
        document.add_paragraph()  # Spacer

        # Create table
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Set column widths
        widths = [Cm(2.5), Cm(11), Cm(5), Cm(3)]
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # Add headers
        headers = table.rows[0].cells
        headers[0].text = 'S. No.'
        headers[1].text = 'Contents'
        headers[2].text = 'Page No.'
        headers[3].text = 'Faculty Signature'
        
        # Format headers
        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(14)
                    run.bold = True
                    run.font.name = "Times New Roman"

        # Add all sections
        all_sections = ordered_sections + [
            s for s in custom_sections 
            if s not in ordered_sections
        ]
        
        for idx, section in enumerate(all_sections, 1):
            display_text = self.section_map.get(section, section.capitalize())
            page_range = self._find_best_page_match(display_text, heading_ranges)
            
            # Format page range string
            if page_range:
                start, end = page_range
                start += 3
                end += 3
                page_str = f"{start}" if start == end else f"{start}-{end}"
            else:
                page_str = ""
            
            # Add row
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = display_text
            row[2].text = page_str
            
            # Format cells
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.name = "Times New Roman"

        # Add empty signature row
        table.add_row().cells

    def _find_best_page_match(
        self, 
        display_text: str, 
        heading_ranges: Dict[str, Tuple[int, int]]
    ) -> Optional[Tuple[int, int]]:
        """
        Find the best matching page range for a section
        Returns: Tuple of (start_page, end_page) or None
        """
        if not heading_ranges:
            return None
            
        # Normalize the display text for matching
        normalized_display = display_text.lower().strip()
        
        # Try exact match first
        for heading, pages in heading_ranges.items():
            if normalized_display == heading.lower():
                return pages
        
        # Try partial matches with common variations
        variations = [
            normalized_display,
            normalized_display.replace(" of the system", ""),
            normalized_display.replace(" of the project", ""),
            normalized_display.replace("study", "").strip(),
            " ".join([word for word in normalized_display.split() if word not in ["the", "of", "a"]])
        ]
        
        for variation in variations:
            for heading, pages in heading_ranges.items():
                if variation and variation in heading.lower():
                    return pages
        
        # Try to find the closest match by section order
        for heading, pages in heading_ranges.items():
            if normalized_display in heading.lower() or heading.lower() in normalized_display:
                return pages
                
        return None