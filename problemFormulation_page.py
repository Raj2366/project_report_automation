import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}

def fetch_problem_formulation(topic, report_type='college'):
    """
    Fetch Problem Formulation content for the given topic using the Gemini API.
    The response will be formatted into a structured document with headings and subheadings.
    """
    # Define the prompt with explicit formatting instructions
    if report_type == 'school':
        prompt = (
            f"Provide a problem formulation for the topic: {topic}. "
            "The problem formulation should include:\n\n"
            "1. A general introduction paragraph explaining the problem specification.\n"
            "2. A list of specific problems or challenges, each with a subheading and a brief description.\n"
            "Format the response as follows:\n\n"
            
            "[Introduction paragraph]\n\n"
            "1. [Subheading 1]: [Description of problem 1]\n\n"
            "2. [Subheading 2]: [Description of problem 2]\n\n"
            "3. [Subheading 3]: [Description of problem 3]\n\n"
            "4. [Subheading 4]: [Description of problem 4]\n\n"
            "5. [Subheading 5]: [Description of problem 5]\n\n"

        )
    else:
        prompt = (
            f"Provide a problem formulation for the topic: {topic}. "
            "The problem formulation should include:\n\n"
            "1. A general introduction paragraph explaining the problem specification.\n"
            "2. A list of specific problems or challenges, each with a subheading and a brief description.\n"
            "Format the response as follows:\n\n"
            
            "[Introduction paragraph]\n\n"
            "1. [Subheading 1]: [Description of problem 1]\n\n"
            "2. [Subheading 2]: [Description of problem 2]\n\n"
            "3. [Subheading 3]: [Description of problem 3]\n\n"
            "4. [Subheading 4]: [Description of problem 4]\n\n"
            "5. [Subheading 5]: [Description of problem 5]\n\n"
        )

    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }
    
    params = {'key': GEMINI_API_KEY}

    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()
        if 'candidates' in content and len(content['candidates']) > 0:
            generated_text = content['candidates'][0]['content']['parts'][0]['text']
            
            # Post-process the generated text to ensure proper formatting
            return generated_text
        else:
            return f"No content found for Problem Formulation on topic: {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching Problem Formulation content: {e}"


def post_process_problem_formulation_to_doc(text, document, report_type='college'):
    """
    Post-process the generated problem formulation text and add it to the Word document.
    """
    # ===== 1. Set Times New Roman as Default Font =====
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # ===== 2. Process Text Line-by-Line =====
    text = text.replace("**", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    intro_processed = False
    previous_line = None  # Track the last line added to detect duplicates

    if report_type == 'school':
        for line in lines:
            para = document.add_paragraph(line)
            para.paragraph_format.space_after = Pt(12)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                
    else:
        intro_processed = False
        previous_line = None

        for line in lines:
            if line == previous_line:
                continue

            if not intro_processed:
                # ===== INTRO PARAGRAPH =====
                para = document.add_paragraph(line)
                para.paragraph_format.space_after = Pt(21)  # Extra space after intro
                intro_processed = True
                previous_line = line  # Store the last added line

                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue

            # === Case 1: Subheading (e.g., "Technical Feasibility: Description")
            if ":" in line and not line.startswith("Feasibility Study"):
                heading, desc = line.split(":", 1)
                heading = heading.strip()
                desc = desc.strip()

                # -- Add Heading (Bold + Underlined) --
                para = document.add_paragraph()
                run_heading = para.add_run(f"{heading}:")
                run_heading.bold = True
                run_heading.underline = True
                run_heading.font.name = 'Times New Roman'
                run_heading.font.size = Pt(12)
                
                # -- Add Description (Normal) --
                run_desc = para.add_run(f" {desc}")
                run_desc.font.name = 'Times New Roman'
                run_desc.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # -- Set Spacing After Paragraph --
                # para.paragraph.alignment = 3
                para.paragraph_format.space_after = Pt(19)  # 12pt spacing after subheadings

            # === Case 2: Normal Paragraph (e.g., Introduction)
            else:
                para = document.add_paragraph(line)
                para.paragraph_format.space_after = Pt(12)  # 6pt spacing after normal paragraphs
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Ensure all runs use Times New Roman
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
