import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Replace 'YOUR_VALID_API_KEY' with your actual Gemini API key.
GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}

def fetch_feasibility_study(topic, report_type='college'):
    """
    Fetch Feasibility Study content for the given topic using the Gemini API.
    The response will be formatted into a structured document with headings and subheadings.
    """
    # Define the prompt with explicit formatting instructions
    if report_type == 'school':
        prompt = (
            f"Provide a feasibility study for the topic: {topic}. "
            "The feasibility study should include:\n\n"
            "1. A general introduction paragraph explaining the feasibility study.\n"
            "2. A list of key components, each with a subheading and a brief description.\n"
            "Format the response as follows:\n\n"

            "[Introduction paragraph]\n\n"
            "1.[Technical Feasibility]: [Description of feasibility 1]\n\n"
            "2.[Operational Feasibility]: [Description of feasibility 2]\n\n"
            "3.[Financial Feasibility]: [Description of feasibility 3]\n\n"
        )
    else:
        prompt = (
            f"Provide a feasibility study for the topic: {topic}. "
            "The feasibility study should include:\n\n"
            "1. A general introduction paragraph explaining the feasibility study.\n"
            "2. A list of key components, each with a subheading and a brief description.\n"
            "Format the response as follows:\n\n"

            "[Introduction paragraph]\n\n"
            "1.[Technical Feasibility]: [Description of feasibility 1]\n\n"
            "2.[Operational Feasibility]: [Description of feasibility 2]\n\n"
            "3.[Financial Feasibility]: [Description of feasibility 3]\n\n"
        )

    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }
    
    params = {'key': GEMINI_API_KEY}

    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()
        if 'candidates' in content and len(content['candidates']) > 0:
            generated_text = content['candidates'][0]['content']['parts'][0]['text']
            return generated_text
        else:
            return f"No content found for Feasibility Study on topic: {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching Feasibility Study content: {e}"

def post_process_feasibility_study_to_doc(text, document, report_type='college'):
    """
    Post-process the generated feasibility study text and add it to the Word document.
    """
   # ===== 1. Set Times New Roman as Default Font =====
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # ===== 2. Process Text Line-by-Line =====
    text = text.replace("**", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    if report_type == 'school':
        for line in lines:
            para = document.add_paragraph(line)
            # para.paragraph_format.space_after = Pt(12)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Improved justification control
            para.paragraph_format.widow_control = True  # Prevent single words on last line
            para.paragraph_format.keep_together = True  # Keep paragraph together
            para.paragraph_format.keep_with_next = True
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            # Adjust hyphenation (helps with justification)
            para.paragraph_format.word_wrap = True
            para.paragraph_format.hyphenation = True

    else:
        intro_processed = False
        previous_line = None  # Track the last line added to detect duplicates

        for line in lines:
            # Skip if current line is identical to the previous line (duplicate)
            if line == previous_line:
                continue

            if not intro_processed:
                # ===== INTRO PARAGRAPH =====
                para = document.add_paragraph(line)
                # para.paragraph_format.space_after = Pt(21)  # Extra space after intro
                intro_processed = True
                previous_line = line  # Store the last added line

                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue

            # === Case 1: Subheading (e.g., "Technical Feasibility: Description")
            if ":" in line and not line.startswith("Feasibility Study"):
                heading, desc = line.split(":", 1)
                heading = heading.strip()
                desc = desc.strip()

                # -- Add Heading (Bold + Underlined) --
                para = document.add_paragraph()
                run_heading = para.add_run(f"{heading}:")
                run_heading.bold = True
                run_heading.underline = True
                run_heading.font.name = 'Times New Roman'
                run_heading.font.size = Pt(12)
                
                # -- Add Description (Normal) --
                run_desc = para.add_run(f" {desc}")
                run_desc.font.name = 'Times New Roman'
                run_desc.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True  # Prevent single words on last line
                para.paragraph_format.keep_together = True  # Keep paragraph together
                para.paragraph_format.keep_with_next = True

                # -- Set Spacing After Paragraph --
                # para.paragraph_format.space_after = Pt(1)  # 12pt spacing after subheadings

                # Adjust hyphenation (helps with justification)
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True

            # === Case 2: Normal Paragraph (e.g., Introduction)
            else:
                para = document.add_paragraph(line)
                # para.paragraph_format.space_after = Pt(12)  # 6pt spacing after normal paragraphs
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Improved justification control
                para.paragraph_format.widow_control = True  # Prevent single words on last line
                para.paragraph_format.keep_together = True  # Keep paragraph together
                para.paragraph_format.keep_with_next = True
                
                # Ensure all runs use Times New Roman
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

                # Adjust hyphenation (helps with justification)
                para.paragraph_format.word_wrap = True
                para.paragraph_format.hyphenation = True


                

