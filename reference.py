import requests
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import docx.oxml.shared
import docx.opc.constants

# Replace 'YOUR_VALID_API_KEY' with your actual Gemini API key.
GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}

def fetch_references(topic, report_type='college'):
    """
    Fetch References content for the given topic using the Gemini API.
    The response will be formatted with numbered references and clickable links.
    """
    if report_type == 'school':
        prompt = (
            f"Generate reference sources for a {topic} project with this format:\n\n"
            "1. Websites:\n"
            "[Website URL 1]\n"
            "[Website URL 2]\n"
            "[Website URL 3]\n\n"
            "2. YouTube:\n"
            "[YouTube URL 1]\n\n"
            "3. Books:\n"
            "[Book Title] by [Author], [Year]\n\n"
            "Example for Hospital Management System:\n"
            "1. Websites:\n"
            "https://www.ncbi.nlm.nih.gov/pmc/articles/PMC7122418/\n"
            "https://www.sciencedirect.com/science/article/pii/S1386505618305983\n\n"
            "2. YouTube:\n"
            "https://www.youtube.com/watch?v=9Hhqj6FVBQ8\n\n"
            "3. Books:\n"
            "Hospital Information Systems by John Smith, 2020\n\n"
            f"Now create references for {topic}:"
        )
    else:
        prompt = (
            f"Generate reference sources for a {topic} project with this format:\n\n"
            "1. Websites:\n"
            "[Website URL 1]\n"
            "[Website URL 2]\n"
            "[Website URL 3]\n\n"
            "2. YouTube:\n"
            "[YouTube URL 1]\n\n"
            "3. Books:\n"
            "[Book Title] by [Author], [Year]\n\n"
            "Example for Hospital Management System:\n"
            "1. Websites:\n"
            "https://www.ncbi.nlm.nih.gov/pmc/articles/PMC7122418/\n"
            "https://www.sciencedirect.com/science/article/pii/S1386505618305983\n\n"
            "2. YouTube:\n"
            "https://www.youtube.com/watch?v=9Hhqj6FVBQ8\n\n"
            "3. Books:\n"
            "Hospital Information Systems by John Smith, 2020\n\n"
            f"Now create references for {topic}:"
        )

    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }
    
    params = {'key': GEMINI_API_KEY}

    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()
        
        if 'candidates' in content and len(content['candidates']) > 0:
            return content['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"No references generated for {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching References content: {e}"

def post_process_references_to_doc(text, document, report_type='college'):
    """
    Post-process the generated references text and add it to the Word document.
    Creates clickable hyperlinks for URLs with proper formatting.
    """
    # Set Times New Roman as Default Font
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


    # Clean and split text
    text = text.replace("**", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    current_category = None

    for line in lines:
        # Check if line starts a new category (number followed by period)
        if line and line[0].isdigit() and '.' in line.split()[0]:
            current_category = line
            para = document.add_paragraph(current_category)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
            
            # Improved justification control for category headings
            para.paragraph_format.widow_control = True
            para.paragraph_format.keep_together = True
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.word_wrap = True
            para.paragraph_format.hyphenation = True
            
        elif line.startswith(('http://', 'https://')):
            # Add hyperlink with proper formatting
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add the URL as a hyperlink
            add_hyperlink(para, line, line, '0000FF', True)
            
            # Formatting for hyperlinks
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            # Improved justification control for links
            para.paragraph_format.widow_control = True
            para.paragraph_format.keep_together = True
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.word_wrap = True
            para.paragraph_format.hyphenation = True
            
        else:
            # Regular text line
            para = document.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            # Improved justification control for regular text
            para.paragraph_format.widow_control = True
            para.paragraph_format.keep_together = True
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.word_wrap = True
            para.paragraph_format.hyphenation = True

def add_hyperlink(paragraph, text, url, color=None, underline=True):
    """
    Add a hyperlink to a paragraph with proper Times New Roman formatting.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, 
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, 
        is_external=True
    )

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Set Times New Roman font
    font = docx.oxml.shared.OxmlElement('w:rFonts')
    font.set(docx.oxml.shared.qn('w:ascii'), 'Times New Roman')
    font.set(docx.oxml.shared.qn('w:hAnsi'), 'Times New Roman')
    font.set(docx.oxml.shared.qn('w:eastAsia'), 'Times New Roman')
    rPr.append(font)

    # Set font size to 12pt
    size = docx.oxml.shared.OxmlElement('w:sz')
    size.set(docx.oxml.shared.qn('w:val'), '24')  # 24 half-points = 12pt
    rPr.append(size)

    if color is not None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink