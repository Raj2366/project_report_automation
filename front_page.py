from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def add_multiple_logos(document, logo_paths):
    """Add logos to the document.
    - 1 logo: center aligned
    - 2+ logos: arranged in pairs per row (left and right aligned)
    """
    if not logo_paths:
        return

    try:
        num_logos = len(logo_paths)

        if num_logos == 1:
            # Center the single logo
            document.add_paragraph()  # spacing
            document.add_paragraph()
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(logo_paths[0], width=Inches(2))
            document.add_paragraph()
            document.add_paragraph()
            return

        # For 2 or more logos – use table layout (2 per row)
        table = document.add_table(rows=0, cols=2)
        table.autofit = True

        i = 0
        while i < num_logos:
            row = table.add_row().cells

            # Left cell
            if os.path.exists(logo_paths[i]):
                p_left = row[0].paragraphs[0]
                p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_left = p_left.add_run()
                run_left.add_picture(logo_paths[i], width=Inches(2))

            # Right cell (only if another logo exists)
            if i + 1 < num_logos and os.path.exists(logo_paths[i + 1]):
                p_right = row[1].paragraphs[0]
                p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_right = p_right.add_run()
                run_right.add_picture(logo_paths[i + 1], width=Inches(2))

            i += 2

        document.add_paragraph()
        document.add_paragraph()

    except Exception as e:
        print(f"Error adding logos to document: {e}")


def create_front_page(document, topic, submitted_by, submitted_to, course_name, semester_name, designate, institute_name, location, enrollment, department, logo_paths=None, **kwargs):
    """Create college report front page with multiple logos"""
    try:
        # Add multiple logos if provided
        if logo_paths:
            add_multiple_logos(document, logo_paths)
    
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(40) 
        run = title.add_run(topic.upper())
        run.bold = True
        run.underline = True
        run.font.size = Pt(21)
        run.font.name = 'Times New Roman'

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("A PROJECT REPORT")
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Times New Roman'

        subtitle.paragraph_format.space_after = Pt(50)
        # Add spacing using an empty paragraph
        # document.add_paragraph()  # This adds one line of space

        text1 = document.add_paragraph()
        text1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = text1.add_run("Submitted in partial fulfilment of the requirements\nfor the award of the degree of")
        run.font.size = Pt(16)
        run.italic = True
        text1.paragraph_format.space_after = Pt(0)  # Remove space after

        degree = document.add_paragraph()
        degree.alignment = WD_ALIGN_PARAGRAPH.CENTER
        degree.paragraph_format.space_before = Pt(0)  # Remove space before
        run = degree.add_run(course_name.upper())
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'
        degree.paragraph_format.space_after = Pt(0)

        semester = document.add_paragraph()
        semester.alignment = WD_ALIGN_PARAGRAPH.CENTER
        semester.paragraph_format.space_before = Pt(0)  # Remove space before
        run = semester.add_run(f"({semester_name.upper()})" if semester_name else "SEMESTER NAME")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True

        # for _ in range(3):
        #     document.add_paragraph()
        semester.paragraph_format.space_after = Pt(80)

        # Create table with 1 row and 2 columns
        table = document.add_table(rows=2, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Submitted By
        cell_left = table.cell(0, 0)
        submitted_by_paragraph = cell_left.paragraphs[0]
        # submitted_by_paragraph = document.add_paragraph()
        submitted_by_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = submitted_by_paragraph.add_run(f"Submitted By: \nMr. {submitted_by}")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True
        submitted_by_paragraph.paragraph_format.space_after = Pt(0)

        # Row 2 - Enrollment No
        # Row 2 - Enrollment No
        cell_2 = table.cell(1, 0)
        enrollment_no = cell_2.paragraphs[0]
        # enrollment_no = document.add_paragraph()
        enrollment_no.alignment = WD_ALIGN_PARAGRAPH.LEFT
        enrollment_no.paragraph_format.space_before = Pt(0)
        run = enrollment_no.add_run(enrollment.upper())
        run.font.size = Pt(13)
        run.bold = True
        run.font.name = 'Times New Roman'

        # document.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Right cell - Submitted To
        cell_right = table.cell(0, 1)
        submitted_to_paragraph = cell_right.paragraphs[0]
        # submitted_to_paragraph = document.add_paragraph()
        submitted_to_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = submitted_to_paragraph.add_run(f"Submitted to: \nMrs. {submitted_to}")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True


        # Designation
        cell_right_2 = table.cell(1, 1)
        designation = cell_right_2.paragraphs[0]
        # designation = document.add_paragraph()
        designation.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = designation.add_run(f"({designate})")
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = 'Times New Roman'

        designation.paragraph_format.space_after = Pt(40)

        university_department = document.add_paragraph()
        university_department.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = university_department.add_run(department.upper())
        run.font.size = Pt(16)
        run.bold = True
        run.font.name = 'Times New Roman'
        university_department.paragraph_format.space_after = Pt(0)

        # Institution
        institution = document.add_paragraph()
        institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
        institution.paragraph_format.space_before = Pt(0)
        run = institution.add_run(institute_name.title())
        run.font.size = Pt(16)
        run.bold = True
        run.font.name = 'Times New Roman'
        institution.paragraph_format.space_after = Pt(0)


        affiliated_university = document.add_paragraph()
        affiliated_university.alignment = WD_ALIGN_PARAGRAPH.CENTER
        affiliated_university.paragraph_format.space_before = Pt(0)
        run = affiliated_university.add_run(location.upper())
        run.font.size = Pt(16)
        run.bold = True
        run.font.name = 'Times New Roman'
    except Exception as e:
        print(f"Error creating front page: {e}")

def create_school_front_page(document, topic, submitted_by, submitted_to, course_name, semester_name, designate, institute_name, location, enrollment, department, logo_paths=None, **kwargs):
    """Create school report front page with multiple logos"""
    try:
        # Add multiple logos if provided
        if logo_paths:
            add_multiple_logos(document, logo_paths)
    
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(40) 
        run = title.add_run(topic.upper())
        run.bold = True
        run.underline = True
        run.font.size = Pt(21)
        run.font.name = 'Times New Roman'

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("A PROJECT SYNOPSIS")
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Times New Roman'

        subtitle.paragraph_format.space_after = Pt(50)
        # Add spacing using an empty paragraph
        # document.add_paragraph()  # This adds one line of space

        text1 = document.add_paragraph()
        text1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = text1.add_run("SUBMITTED IN PARTIAL FULFILMENT OF THE REQUIREMENTS\nFOR THE AWARD OF THE DEGREE OF")
        run.font.size = Pt(14)
        # run.italic = True
        # text1.paragraph_format.space_after = Pt(0)  # Remove space after

        degree = document.add_paragraph()
        degree.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # degree.paragraph_format.space_before = Pt(0)  # Remove space before
        run = degree.add_run(course_name.upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        degree.paragraph_format.space_after = Pt(0)

        semester = document.add_paragraph()
        semester.alignment = WD_ALIGN_PARAGRAPH.CENTER
        semester.paragraph_format.space_before = Pt(0)  # Remove space before
        run = semester.add_run(f"({semester_name.upper()})" if semester_name else "SEMESTER NAME")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True

        # for _ in range(3):
        #     document.add_paragraph()
        semester.paragraph_format.space_after = Pt(80)

        # Create table with 1 row and 2 columns
        table = document.add_table(rows=2, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Submitted By
        cell_left = table.cell(0, 0)
        submitted_by_paragraph = cell_left.paragraphs[0]
        # submitted_by_paragraph = document.add_paragraph()
        submitted_by_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = submitted_by_paragraph.add_run(f"Submitted By: \nMr. {submitted_by}")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True
        submitted_by_paragraph.paragraph_format.space_after = Pt(0)

        # Row 2 - Enrollment No
        # Row 2 - Enrollment No
        cell_2 = table.cell(1, 0)
        enrollment_no = cell_2.paragraphs[0]
        # enrollment_no = document.add_paragraph()
        enrollment_no.alignment = WD_ALIGN_PARAGRAPH.LEFT
        enrollment_no.paragraph_format.space_before = Pt(0)
        run = enrollment_no.add_run(enrollment.upper())
        run.font.size = Pt(13)
        run.bold = True
        run.font.name = 'Times New Roman'

        # document.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Right cell - Submitted To
        cell_right = table.cell(0, 1)
        submitted_to_paragraph = cell_right.paragraphs[0]
        # submitted_to_paragraph = document.add_paragraph()
        submitted_to_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = submitted_to_paragraph.add_run(f"Submitted to: \nMrs. {submitted_to}")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True

        # Designation
        cell_right_2 = table.cell(1, 1)
        designation = cell_right_2.paragraphs[0]
        # designation = document.add_paragraph()
        designation.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = designation.add_run(f"({designate})")
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = 'Times New Roman'

        designation.paragraph_format.space_after = Pt(40)

        university_department = document.add_paragraph()
        university_department.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = university_department.add_run(department.upper())
        run.font.size = Pt(16)
        run.bold = True
        run.font.name = 'Times New Roman'
        university_department.paragraph_format.space_after = Pt(0)

        # Institution
        institution = document.add_paragraph()
        institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
        institution.paragraph_format.space_before = Pt(0)
        run = institution.add_run(institute_name.upper())
        run.font.size = Pt(13)
        run.bold = True
        run.font.name = 'Times New Roman'
        institution.paragraph_format.space_after = Pt(0)

        affiliated_university = document.add_paragraph()
        affiliated_university.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = affiliated_university.add_run(location.upper())
        run.font.size = Pt(13)
        run.bold = True
        run.font.name = 'Times New Roman'


    except Exception as e:
        print(f"Error creating school front page: {e}")

