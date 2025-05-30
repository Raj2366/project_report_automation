import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}



def fetch_custom_section(topic, report_type='college'):
    """
    Fetch AI-generated content for a custom section.
    """
    topic = topic.strip()
    if report_type == "school":
        prompt = (
        f"Write a structured and well-organized explanation for the topic: {topic}.\n"
        "1. [Main Point 1]: [Detailed explanation of this aspect of the topic]\n"
        "2. [Main Point 2]: [Detailed explanation of this aspect of the topic]\n"
        "3. [Main Point 3]: [Detailed explanation of this aspect of the topic]\n"
        "4. [Main Point 4]: [Detailed explanation of this aspect of the topic]\n"
        "Example for the topic 'Hospital Management System':\n"
        "1. Predictive Analytics Engine: Develop ML models using Python/TensorFlow to predict admission rates and medication needs based on historical data patterns.\n"
        "2. Blockchain Medical Records: Implement Ethereum-based decentralized ledger for tamper-proof records with patient-controlled access permissions.\n"
        "3. IoT Smart Hospital: Network of sensors to monitor equipment status and patient movements in real-time with automated maintenance alerts.\n"
        "4. AR Surgical Assistance: Augmented reality system to overlay patient vitals and imaging data during surgical procedures for precision.\n"
        f"Now write a structured explanation for the topic: {topic}."
        )
    else:
        prompt = (
        f"Write a structured and well-organized explanation for the topic: {topic}.\n"
        "1. [Main Point 1]: [Detailed explanation of this aspect of the topic]\n"
        "2. [Main Point 2]: [Detailed explanation of this aspect of the topic]\n"
        "3. [Main Point 3]: [Detailed explanation of this aspect of the topic]\n"
        "4. [Main Point 4]: [Detailed explanation of this aspect of the topic]\n"
        "Example for the topic 'Hospital Management System':\n"
        "1. Predictive Analytics Engine: Develop ML models using Python/TensorFlow to predict admission rates and medication needs based on historical data patterns.\n"
        "2. Blockchain Medical Records: Implement Ethereum-based decentralized ledger for tamper-proof records with patient-controlled access permissions.\n"
        "3. IoT Smart Hospital: Network of sensors to monitor equipment status and patient movements in real-time with automated maintenance alerts.\n"
        "4. AR Surgical Assistance: Augmented reality system to overlay patient vitals and imaging data during surgical procedures for precision.\n"
        f"Now write a structured explanation for the topic: {topic}."
        )

    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }

    params = {'key': GEMINI_API_KEY}
    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()
        if 'candidates' in content and len(content['candidates']) > 0:
            return content['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"No content found for {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching content: {e}"
    

def post_process_custom_to_doc(text, document, report_type='college'):
    """
    Post-process the generated objective text and add it to the Word document.
    Applies different formatting for school vs college reports.
    """
    # Set Times New Roman as Default Font
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Clean and split text
    text = text.replace("**", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    def format_paragraph(paragraph, is_heading=False):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if is_heading:
                run.bold = True
                run.underline = True

        paragraph.paragraph_format.widow_control = True
        paragraph.paragraph_format.keep_together = True

        if is_heading:
            paragraph.paragraph_format.keep_with_next = True
        else:
            paragraph.paragraph_format.keep_with_next = False

        paragraph.paragraph_format.word_wrap = True
        paragraph.paragraph_format.hyphenation = True

    if report_type == 'school':
        for line in lines:
            is_heading = line[0].isdigit() and '.' in line.split()[0]
            para = document.add_paragraph(line)
            format_paragraph(para, is_heading=is_heading)
    else:
        previous_line = None

        for line in lines:
            if line == previous_line:
                continue

            is_heading = line[0].isdigit() and '.' in line.split()[0]

            if ":" in line and is_heading:
                heading, desc = line.split(":", 1)
                para = document.add_paragraph()
                
                run_heading = para.add_run(f"{heading.strip()}:")
                run_heading.bold = True
                run_heading.underline = True
                run_heading.font.size = Pt(12)

                run_desc = para.add_run(f" {desc.strip()}")
                run_desc.font.size = Pt(12)

                format_paragraph(para)
            else:
                para = document.add_paragraph(line)
                format_paragraph(para, is_heading=is_heading)

            previous_line = line



    

