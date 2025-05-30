from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement, parse_xml
import os
import tempfile
import shutil
import json
import requests
import uuid
import logging
import mimetypes
from werkzeug.utils import secure_filename
from PIL import Image
from introduction_page import fetch_introduction, post_process_introduction_to_doc
from objective_page import fetch_objective, post_process_objective_to_doc
from future_scope_page import fetch_future_scope, post_process_future_scope_to_doc
from conclusion_page import fetch_conclusion, post_process_conclusion
from content_api import fetch_custom_section, post_process_custom_to_doc
from front_page import create_front_page, create_school_front_page
from dfd_generator import generate_dfd_text, generate_dfd_image
from flowchart_page import generate_flowchart_text, generate_flowchart_image
from problemFormulation_page import fetch_problem_formulation, post_process_problem_formulation_to_doc
from feasibility_page import fetch_feasibility_study, post_process_feasibility_study_to_doc
from unique_features import fetch_unique_features, post_process_unique_features_to_doc
from abstract_page import fetch_abstract, post_process_abstract
from acknowledgement import generate_acknowledgement_page
from modules_desc import fetch_modules_and_descriptions, post_process_modules_to_doc
from limitations import fetch_limitations, post_process_limitations_to_doc
from reference import fetch_references, post_process_references_to_doc, add_hyperlink
from index_page import IndexGenerator
from image_page import add_image_page
from image_generate import generate_image

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def add_heading(document, text, size=18, bold=True, underline=True):
    """Add a heading to the document with custom formatting"""
    heading = document.add_paragraph()
    run = heading.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.font.name = "Times New Roman"
    heading.alignment = 1  # Center alignment

def save_uploaded_files(uploaded_files, temp_dir):
    """Save uploaded files to temp location"""
    saved_paths = []
    
    for file in uploaded_files:
        if file.filename == '':
            continue
            
        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(temp_dir, filename)
            file.save(filepath)
            saved_paths.append(filepath)
        except Exception as e:
            print(f"Error saving file {file.filename}: {e}")
    
    return saved_paths

def process_images(image_files, temp_dir):
    """Process both uploaded and generated images"""
    processed_images = []
    
    for img in image_files:
        if isinstance(img, str):  # This is a path to a generated image
            try:
                if os.path.exists(img):
                    filename = os.path.basename(img)
                    dest_path = os.path.join(temp_dir, filename)
                    shutil.copy(img, dest_path)
                    processed_images.append(dest_path)
            except Exception as e:
                print(f"Error copying generated image {img}: {e}")
        else:
            try:
                filename = secure_filename(img.filename)
                filepath = os.path.join(temp_dir, filename)
                img.save(filepath)
                processed_images.append(filepath)
            except Exception as e:
                print(f"Error saving uploaded image {img.filename}: {e}")
    
    return processed_images

def get_image_extension(content_type):
    """Get file extension from content type"""
    return mimetypes.guess_extension(content_type)

def add_logos_to_document(document, logo_paths):
    """Add multiple logos to the document from saved paths"""
    if not logo_paths:
        return
        
    try:
        # Add space before logos
        document.add_paragraph()
        
        # Create centered paragraph for logos
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = 1  # Center alignment
        
        # Add each logo
        for logo_path in logo_paths:
            try:
                if os.path.exists(logo_path):
                    run = logo_paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(1.5))
                    run.add_break()  # Add space between logos
            except Exception as e:
                print(f"Error adding logo {logo_path}: {e}")
                continue
        
        # Add space after logos
        document.add_paragraph()
        document.add_paragraph()
    except Exception as e:
        print(f"Error adding logos to document: {e}")

def remove_blank_pages(document):
    """Remove blank pages from the document"""
    for paragraph in document.paragraphs:
        if len(paragraph.text.strip()) == 0 and len(paragraph.runs) == 0:
            p = paragraph._element
            p.getparent().remove(p)

def add_page_numbers(document):
    """Add page numbers to the footer"""
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = 1
        
        run = paragraph.add_run()
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char)

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'PAGE'
        run._r.append(instr_text)

        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char)

def set_page_formatting(document, report_type='college'):
    """Set margins and page borders"""
    section = document.sections[-1]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    if report_type in ['college', 'school']:
        sectPr = section._sectPr
        borders = parse_xml(
            r'''
            <w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                         xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">
                <w:top w:val="double" w:sz="6" w:space="28" w:color="000000"/>
                <w:left w:val="double" w:sz="6" w:space="28" w:color="000000"/>
                <w:bottom w:val="double" w:sz="6" w:space="32" w:color="000000"/>
                <w:right w:val="double" w:sz="6" w:space="28" w:color="000000"/>
            </w:pgBorders>
            '''
        )
        sectPr.append(borders)

def create_college_report(document, topic, selected_sections, custom_sections, ordered_sections, project_images=None, code_images=None, temp_dir=None, **kwargs):
    """Generate a college-level project report"""
    section_processors = {
        "Introduction": lambda: fetch_introduction(topic),
        "Objective": lambda: fetch_objective(topic),
        "Future Scope": lambda: fetch_future_scope(topic),
        "Conclusion": lambda: fetch_conclusion(topic),
        "DFD": lambda: generate_dfd_image(generate_dfd_text(topic), topic),
        "Flowchart": lambda: generate_flowchart_image(generate_flowchart_text(topic), topic),
        "Problem Formulation": lambda: fetch_problem_formulation(topic),
        "Feasibility": lambda: fetch_feasibility_study(topic),
        "Unique Features": lambda: fetch_unique_features(topic),
        "Future Scope": lambda: fetch_future_scope(topic),
        "Reference": lambda: fetch_references(topic),
        "Project Images": lambda: project_images,
        "Code Screenshots": lambda: code_images
    }

    for section_name in ordered_sections:
        section_name = section_name.strip()
        
        # Skip empty sections
        if not section_name:
            continue
            
        # Add page break for each new section except the first one
        if document.paragraphs:
            document.add_page_break()
        
        # Handle custom sections
        if section_name in custom_sections:
            add_heading(document,f"{section_name.title()}", size=16, bold=True, underline=True)
            content = fetch_custom_section(section_name)
            format_custom = post_process_custom_to_doc(content, document)
            document.add_paragraph(format_custom)
            continue
            
        # Handle predefined sections
        if section_name in section_processors:
            content = section_processors[section_name]()
            
            if section_name == "Introduction":
                add_heading(document, f' "_____{topic.upper()}_____" ', size=18, bold=True)
                add_heading(document, "INTRODUCTION\n", size=16, bold=True, underline=True)
                post_process_introduction_to_doc(content, document)
            elif section_name == "Objective":
                add_heading(document, "Objective\n", size=16, bold=True, underline=True)
                post_process_objective_to_doc(content, document)
            elif section_name == "Unique Features":
                add_heading(document, "Unique Features of the System\n", size=16, bold=True, underline=True)
                post_process_unique_features_to_doc(content, document)
            elif section_name == "Problem Formulation":
                add_heading(document, "Problem Formulation\n", size=16, bold=True, underline=True)
                post_process_problem_formulation_to_doc(content, document)
            elif section_name == "Feasibility":
                add_heading(document, "REQUIREMENT ANALYSIS AND SYSTEM SPECIFICATION\n", size=16, bold=True, underline=True)
                add_heading(document, "Feasibility Study of Project\n", size=16, bold=True, underline=True)
                post_process_feasibility_study_to_doc(content, document)
            elif section_name == "Conclusion":
                add_heading(document, "Conclusion\n", size=16, bold=True, underline=True)
                post_process_conclusion(content, document)
            elif section_name == "Future Scope":
                add_heading(document, "Future Scope\n", size=16, bold=True, underline=True)
                post_process_future_scope_to_doc(content, document)
            elif section_name == "DFD":
                # document.add_paragraph("Data Flow Diagram:")
                paragraph = document.add_paragraph()
                run = paragraph.add_run("Data Flow Diagram:")
                run.bold = True
                if isinstance(content, str) and content.startswith("Error"):
                    document.add_paragraph(content)
                else:
                    try:
                        document.add_picture(content, width=Inches(6))
                    except Exception as e:
                        document.add_paragraph("Failed to insert DFD image.")
            elif section_name == "Flowchart":
                # document.add_paragraph("Flowchart:")
                paragraph = document.add_paragraph()
                run = paragraph.add_run("Flowchart:")
                run.bold = True
                if isinstance(content, str) and content.startswith("Error"):
                    document.add_paragraph(content)
                else:
                    try:
                        document.add_picture(content, width=Inches(6))
                    except Exception as e:
                        document.add_paragraph("Failed to insert Flowchart image.")
            elif section_name == "Project Images":
                if content:
                    processed_images = process_images(content, temp_dir)
                    # add_image_page(document, processed_images, "Project Snapshots")
                    add_image_page(document, project_images, title="Project Snapshots", caption_prefix="Output")
            elif section_name == "Code Screenshots":
                if content:
                    processed_images = process_images(content, temp_dir)
                    # add_image_page(document, processed_images, "Code Snapshots")
                    add_image_page(document, code_images, title="Code Snapshots", caption_prefix="Program")
            elif section_name == "Reference":
                add_heading(document, "References/Bibliography\n", size=16, bold=True, underline=True)
                post_process_references_to_doc(content, document)
                # add_hyperlink(content, url)
            else:
                add_heading(document,f"{section_name.capitalize()}", size=16, bold=True, underline=True)
                content = fetch_custom_section(section_name)
                format_custom = post_process_custom_to_doc(content, document)
                document.add_paragraph(format_custom)

def create_school_report(document, topic, selected_sections, custom_sections, ordered_sections, project_images=None, temp_dir=None, **kwargs):
    """Generate a school-level project report"""
    section_processors = {
        "Introduction": lambda: fetch_introduction(topic, report_type='school'),
        "Objective": lambda: fetch_objective(topic, report_type='school'),
        "Unique Features": lambda: fetch_unique_features(topic, report_type = 'school'),
        "Feasibility": lambda: fetch_feasibility_study(topic, report_type = 'school'),
        "Modules and Descriptions": lambda:fetch_modules_and_descriptions(topic, report_type = 'school'),
        "Future Scope": lambda: fetch_future_scope(topic, report_type='school'),
        "DFD": lambda: generate_dfd_image(generate_dfd_text(topic), topic),
        "Flowchart": lambda: generate_flowchart_image(generate_flowchart_text(topic), topic),
        "Conclusion": lambda: fetch_conclusion(topic, report_type='school'),
        "Limitations": lambda: fetch_limitations(topic, report_type = 'school'),
        "Reference": lambda: fetch_references(topic, report_type='school'),
        "Project Images": lambda: project_images
    }

    for section_name in ordered_sections:
        section_name = section_name.strip()
        
        if not section_name:
            continue
            
        if document.paragraphs:
            document.add_page_break()
        
        if section_name in custom_sections:
            add_heading(document,f"{section_name.title()}", size=16, bold=True, underline=True)
            content = fetch_custom_section(section_name)
            format_custom = post_process_custom_to_doc(content, document)
            document.add_paragraph(format_custom)
            continue
            
        if section_name in section_processors:
            content = section_processors[section_name]()
            
            if section_name == "Introduction":
                add_heading(document, f' "_____{topic.upper()}_____" ', size=18, bold=True, underline=True)
                add_heading(document, "INTRODUCTION\n", size=16, bold=True, underline=True)
                post_process_introduction_to_doc(content, document)
            elif section_name == "Objective":
                add_heading(document, "Objective\n", size=16, bold=True, underline=True)
                post_process_objective_to_doc(content, document)
            elif section_name == "Unique Features":
                add_heading(document, "Unique Features of the System\n", size=16, bold=True, underline=True)
                post_process_unique_features_to_doc(content, document)
            elif section_name == "Feasibility":
                add_heading(document, "Feasibility Study of Project\n", size=16, bold=True, underline=True)
                post_process_feasibility_study_to_doc(content, document)
            elif section_name == "Modules and Descriptions":
                add_heading(document, "Modules and Descriptions\n", size=16, bold=True, underline=True)
                post_process_modules_to_doc(content, document)
            elif section_name == "Conclusion":
                add_heading(document, "Conclusion\n", size=16, bold=True, underline=True)
                post_process_conclusion(content, document)
            elif section_name == "Future Scope":
                add_heading(document, "Future Scope\n", size=16, bold=True, underline=True)
                post_process_future_scope_to_doc(content, document)
            elif section_name == "Limitations":
                add_heading(document, f"Limitations of the {topic}\n", size=16, bold=True, underline=True)
                post_process_limitations_to_doc(content, document)
            elif section_name == "DFD":
                paragraph = document.add_paragraph()
                run = paragraph.add_run("Data Flow Diagram:")
                run.bold = True
                if isinstance(content, str) and content.startswith("Error"):
                    document.add_paragraph(content)
                else:
                    try:
                        document.add_picture(content, width=Inches(6))
                    except Exception as e:
                        document.add_paragraph("Failed to insert DFD image.")
            elif section_name == "Flowchart":
                paragraph = document.add_paragraph()
                run = paragraph.add_run("Flowchart:")
                run.bold = True
                if isinstance(content, str) and content.startswith("Error"):
                    document.add_paragraph(content)
                else:
                    try:
                        document.add_picture(content, width=Inches(6))
                    except Exception as e:
                        document.add_paragraph("Failed to insert Flowchart image.")
            elif section_name == "Reference":
                add_heading(document, "References/Bibliography\n", size=16, bold=True, underline=True)
                post_process_references_to_doc(content, document)
            elif section_name == "Project Images":
                if content:
                    processed_images = process_images(content, temp_dir)
                    add_image_page(document, project_images, title="Project Snapshots", caption_prefix="Output")
            else:
                document.add_heading(section_name, level=2)
                document.add_paragraph(content)


def create_project_file(topic, submitted_by, submitted_to, selected_sections, custom_sections, ordered_sections, project_images=None, code_images=None, report_type='college', **kwargs):
    # Create output folder
    OUTPUT_FOLDER = "generated_docs/project_reports" if report_type == 'college' else "generated_docs/project_synopsis"
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            # FIRST PASS: Create complete document with all sections
            complete_doc_path = os.path.join(temp_dir, "complete.docx")
            complete_doc = Document()
            
            # Add all content to complete document
            if report_type == 'college':
                create_front_page(complete_doc, topic, submitted_by, submitted_to, **kwargs)
                create_college_report(complete_doc, topic, selected_sections, custom_sections, ordered_sections,
                                    project_images, code_images, temp_dir, **kwargs)
            else:
                create_school_front_page(complete_doc, topic, submitted_by, submitted_to, **kwargs)
                create_school_report(complete_doc, topic, selected_sections, custom_sections, ordered_sections,
                                   project_images, temp_dir, **kwargs)
            
            complete_doc.save(complete_doc_path)

            # SECOND PASS: Create final document with proper ordering
            final_doc = Document()
            
            # Add front matter
            if report_type == 'college':
                create_front_page(final_doc, topic, submitted_by, submitted_to, **kwargs)
            else:
                create_school_front_page(final_doc, topic, submitted_by, submitted_to, **kwargs)
            
            set_page_formatting(final_doc, report_type)

            # Add abstract
            final_doc.add_page_break()
            add_heading(final_doc, "ABSTRACT", size=16, bold=True, underline=True)
            abstract_content = fetch_abstract(topic)
            processed_abstract = post_process_abstract(abstract_content, final_doc)
            final_doc.add_paragraph(processed_abstract)

            final_doc.add_page_break()
            add_heading(final_doc, "ACKNOWLEDGEMENT", size=16, bold=True, underline=True)
            generate_acknowledgement_page(final_doc, submitted_by, submitted_to, **kwargs)
            
            # Add index with reference to complete document
            generator = IndexGenerator()
            generator.create_index_page(final_doc, ordered_sections, custom_sections, complete_doc_path)

            # Add main content to final document
            if report_type == 'college':
                create_college_report(final_doc, topic, selected_sections, custom_sections, ordered_sections,
                                    project_images, code_images, temp_dir, **kwargs)
            else:
                create_school_report(final_doc, topic, selected_sections, custom_sections, ordered_sections,
                                   project_images, temp_dir, **kwargs)

            # Finalize document
            remove_blank_pages(final_doc)
            add_page_numbers(final_doc)

            # Save final document
            filename = f"{topic.replace(' ', '_')}_Report.docx" if report_type == 'college' else f"{topic.replace(' ', '_')}_Synopsis.docx"
            filepath = os.path.join(OUTPUT_FOLDER, filename)
            final_doc.save(filepath)

            return filename
            
        except Exception as e:
            logger.error(f"Error generating project file: {e}")
            raise   

