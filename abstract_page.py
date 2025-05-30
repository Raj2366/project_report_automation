import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


GEMINI_API_KEY = "AIzaSyCmFMR-XyUb2GdR1Hdub5_g_C6xjnpTqaA"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
HEADERS = {"Content-Type": "application/json"}

def fetch_abstract(topic, report_type='college'):
    """
    Fetch Introduction content for the given topic using the Gemini API.
    The introduction will be tailored for either school or college reports.
    """
    # Different prompts for school vs college reports
    if report_type == 'school':
        prompt = (f"Generate a professional 2-3 paragraph abstract about {topic}.\n"
                 "Paragraphs should be separated by exactly one blank line.\n"
                 "First paragraph: Problem statement\n"
                 "Second paragraph: Methodology\n"
                 "Third paragraph (optional): Key findings")
    else:
        prompt = (f"Generate a professional 2-3 paragraph abstract about {topic}.\n"
                 "Paragraphs should be separated by exactly one blank line.\n"
                 "First paragraph: Problem statement\n"
                 "Second paragraph: Methodology\n"
                 "Third paragraph (optional): Key findings")
    
    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }
    
    params = {'key': GEMINI_API_KEY}

    try:
        response = requests.post(GEMINI_URL, headers=HEADERS, params=params, json=data)
        response.raise_for_status()
        content = response.json()
        if 'candidates' in content and len(content['candidates']) > 0:
            return content['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"No content found for Introduction on topic: {topic}."
    except requests.exceptions.RequestException as e:
        return f"Error fetching Introduction content: {e}"


def post_process_abstract(text, document, report_type='college'):
    """
    Format the abstract section and add it to the Word document with proper justification.
    Adds a heading followed by the text with improved paragraph formatting.
    """
    # Set Times New Roman as Default Font
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Process the abstract text
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    
    for para in paragraphs:
        # Create paragraph with optimized justification settings
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        
        # Improved justification control
        p.paragraph_format.widow_control = True  # Prevent single words on last line
        p.paragraph_format.keep_together = True  # Keep paragraph together
        p.paragraph_format.keep_with_next = True
        
        # Set font properties
        run = p.add_run(para)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Adjust hyphenation (helps with justification)
        p.paragraph_format.word_wrap = True
        p.paragraph_format.hyphenation = True

    # Add spacing after abstract section
    document.add_paragraph()